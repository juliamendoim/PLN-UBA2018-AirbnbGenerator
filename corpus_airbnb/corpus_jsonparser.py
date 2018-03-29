import json
import fnmatch
import os
import pandas as pd
import codecs
from docx import Document
 
datos_path = "/Users/juliamilanese/Desktop/corpus_airbnb/reviews/"

def detect_lan(text):
	if text=='en':
		return('Ingles')
	elif text=='pt':
		return('Portugues')
	elif text=='es':
		return('Español')
	elif text=='fr':
		return('Frances')
	elif text=='it':
		return('Italiano')
	else:
		return(text)

# for file in os.walk(datos_path):
# 	for filename in file:
# 		archivo = filename
# 		print(archivo)


 
 #Loopeo por cada archivo
for file in os.listdir(datos_path):
	data_file = open(datos_path+file)
	#data_file = codecs.open(datos_path+file, 'r', 'utf-8')
	data = json.load(data_file)

	#data es un docuumento con 7 comentarios (adentro de reviews)
	i = 0

	documento = Document()
	
	for elemento in data["reviews"]:

		x = (data["reviews"][i]["comments"])
		y = (data["reviews"][i]["language"])
		#z = (data["reviews"][i]["id"])
		
		if y == 'es':
			with codecs.open('/Users/juliamilanese/Desktop/corpus_airbnb/txt/'+str(file[0:-5])+str(i)+'.txt','w', 'utf-8') as f:
				f.write(str(x)) # + '\n' + str(y)) + '\n'+ str(z))
		
	# 	documento.add_paragraph(str(x))
	# 	documento.add_paragraph(detect_lan(str(y)))
	# 	#documento.add_paragraph(str(z))
		
		i += 1

	# documento.save(str('reviewsdocx/'+str(file[0:-5])+'.docx'))



#Abre el archivo .json
#

#Extrae cada comentario, lenguaje en el que está el comentario y id del usuario que comentó. Los escribe en un doc (para subir a WDS)

	